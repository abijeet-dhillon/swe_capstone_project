"""
Text Analyzer Module
Analyzes text files (PDF, DOCX, TXT) and extracts statistical metrics
No LLM required - uses document parsers and basic statistics
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any
from collections import Counter
from dataclasses import dataclass, asdict

# PDF parsing (optional)
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfpage import PDFPage
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    pdf_extract_text = None
    PDFPage = None

# DOCX parsing
from docx import Document


@dataclass
class TextMetrics:
    """Metrics for a single text file"""
    # File metadata
    file_path: str
    file_name: str
    file_type: str
    file_size_bytes: int
    created_time: str
    modified_time: str
    
    # Content metrics
    page_count: Optional[int]
    word_count: int
    sentence_count: int
    paragraph_count: int
    character_count: int
    
    # Reading metrics
    estimated_reading_time_minutes: float
    
    # Lexical metrics
    lexical_diversity: float
    avg_word_length: float
    avg_sentence_length: float
    
    # Keyword analysis
    top_keywords: List[Tuple[str, int]]
    
    # Document structure (DOCX only)
    heading_count: Optional[int] = None
    heading_breakdown: Optional[Dict[str, int]] = None
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return asdict(self)


class TextAnalyzer:
    """Analyzes text files and extracts metrics"""
    
    # Common English stop words
    STOP_WORDS = {
        'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'i',
        'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do', 'at',
        'this', 'but', 'his', 'by', 'from', 'they', 'we', 'say', 'her', 'she',
        'or', 'an', 'will', 'my', 'one', 'all', 'would', 'there', 'their',
        'what', 'so', 'up', 'out', 'if', 'about', 'who', 'get', 'which', 'go',
        'me', 'when', 'make', 'can', 'like', 'time', 'no', 'just', 'him', 'know',
        'take', 'people', 'into', 'year', 'your', 'good', 'some', 'could', 'them',
        'see', 'other', 'than', 'then', 'now', 'look', 'only', 'come', 'its', 'over',
        'think', 'also', 'back', 'after', 'use', 'two', 'how', 'our', 'work',
        'first', 'well', 'way', 'even', 'new', 'want', 'because', 'any', 'these',
        'give', 'day', 'most', 'us', 'is', 'was', 'are', 'been', 'has', 'had',
        'were', 'said', 'did', 'having', 'may', 'should', 'am', 'being', 'such'
    }
    
    def __init__(self):
        """Initialize the text analyzer"""
        pass
    
    def analyze_file(self, file_path: str) -> TextMetrics:
        """
        Analyze a single text file
        
        Args:
            file_path: Path to the file (PDF, DOCX, or TXT)
            
        Returns:
            TextMetrics object with all metrics
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file metadata
        file_type = self._detect_file_type(file_path)
        file_stats = os.stat(file_path)
        
        # Extract text based on file type
        if file_type == 'pdf':
            text, page_count = self._extract_from_pdf(file_path)
            heading_info = None
        elif file_type == 'docx':
            text, page_count, heading_info = self._extract_from_docx(file_path)
        else:  # txt
            text = self._extract_from_txt(file_path)
            page_count = None
            heading_info = None
        
        # Calculate all metrics
        basic_metrics = self._calculate_basic_metrics(text)
        lexical_metrics = self._calculate_lexical_metrics(text)
        keywords = self._extract_keywords(text, top_n=10)
        reading_time = self._estimate_reading_time(basic_metrics['word_count'])
        
        # Build metrics object
        metrics = TextMetrics(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=file_type,
            file_size_bytes=file_stats.st_size,
            created_time=datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
            modified_time=datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            page_count=page_count,
            word_count=basic_metrics['word_count'],
            sentence_count=basic_metrics['sentence_count'],
            paragraph_count=basic_metrics['paragraph_count'],
            character_count=basic_metrics['character_count'],
            estimated_reading_time_minutes=reading_time,
            lexical_diversity=lexical_metrics['lexical_diversity'],
            avg_word_length=lexical_metrics['avg_word_length'],
            avg_sentence_length=lexical_metrics['avg_sentence_length'],
            top_keywords=keywords,
            heading_count=heading_info.get('heading_count') if heading_info else None,
            heading_breakdown=heading_info.get('heading_breakdown') if heading_info else None
        )
        
        return metrics
    
    def analyze_batch(self, file_paths: List[str]) -> Dict[str, Any]:
        """
        Analyze multiple files and return individual + aggregate metrics
        
        Args:
            file_paths: List of file paths
            
        Returns:
            Dictionary with 'files' (list of individual metrics) and 'totals' (aggregate)
        """
        results = []
        
        for file_path in file_paths:
            try:
                metrics = self.analyze_file(file_path)
                results.append(metrics)
            except Exception as e:
                print(f"Error analyzing {file_path}: {e}")
                continue
        
        # Calculate aggregate metrics
        totals = self._calculate_totals(results)
        
        return {
            'files': [m.to_dict() for m in results],
            'totals': totals
        }
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect file type from extension"""
        ext = file_path.suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'docx'
        elif ext in ['.txt', '.md']:
            return 'txt'
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_from_pdf(self, file_path: Path) -> Tuple[str, int]:
        """Extract text and page count from PDF"""
        try:
            if not PDF_AVAILABLE:
                raise ValueError("PDF parsing not available (pdfminer not installed)")
            text = pdf_extract_text(str(file_path))
            
            # Count pages
            with open(file_path, 'rb') as f:
                page_count = len(list(PDFPage.get_pages(f)))
            
            return text, page_count
        except Exception as e:
            raise ValueError(f"Error reading PDF: {e}")
    
    def _extract_from_docx(self, file_path: Path) -> Tuple[str, Optional[int], Dict]:
        """Extract text and structure from DOCX"""
        try:
            doc = Document(str(file_path))
            
            # Extract text
            text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # Count headings
            heading_counts = {}
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    heading_counts[para.style.name] = heading_counts.get(para.style.name, 0) + 1
            
            heading_info = {
                'heading_count': sum(heading_counts.values()),
                'heading_breakdown': heading_counts if heading_counts else None
            }
            
            # DOCX doesn't have reliable page count without rendering
            return text, None, heading_info
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {e}")
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _calculate_basic_metrics(self, text: str) -> Dict[str, int]:
        """Calculate word, sentence, paragraph counts"""
        # Word count
        words = [w for w in text.split() if w.strip()]
        word_count = len(words)
        
        # Sentence count (simple regex split)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s for s in sentences if s.strip()]
        sentence_count = len(sentences)
        
        # Paragraph count (split on double newlines)
        paragraphs = text.split('\n\n')
        paragraphs = [p for p in paragraphs if p.strip()]
        paragraph_count = len(paragraphs)
        
        # Character count (excluding whitespace)
        char_count = len(text.replace(' ', '').replace('\n', '').replace('\t', ''))
        
        return {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'character_count': char_count
        }
    
    def _calculate_lexical_metrics(self, text: str) -> Dict[str, float]:
        """Calculate lexical diversity and averages"""
        words = [w.lower() for w in text.split() if w.strip()]
        
        if not words:
            return {
                'lexical_diversity': 0.0,
                'avg_word_length': 0.0,
                'avg_sentence_length': 0.0
            }
        
        # Lexical diversity (unique words / total words)
        unique_words = set(words)
        diversity = len(unique_words) / len(words)
        
        # Average word length
        avg_word_len = sum(len(w) for w in words) / len(words)
        
        # Average sentence length
        sentences = re.split(r'[.!?]+', text)
        sentences = [s for s in sentences if s.strip()]
        avg_sent_len = len(words) / len(sentences) if sentences else 0
        
        return {
            'lexical_diversity': round(diversity, 4),
            'avg_word_length': round(avg_word_len, 2),
            'avg_sentence_length': round(avg_sent_len, 2)
        }
    
    def _extract_keywords(self, text: str, top_n: int = 10) -> List[Tuple[str, int]]:
        """Extract top keywords by frequency"""
        # Tokenize: only words with 3+ letters
        words = re.findall(r'\b[a-z]{3,}\b', text.lower())
        
        # Remove stop words
        words = [w for w in words if w not in self.STOP_WORDS]
        
        # Count frequencies
        freq = Counter(words)
        
        return freq.most_common(top_n)
    
    def _estimate_reading_time(self, word_count: int) -> float:
        """Estimate reading time in minutes (200 words/min average)"""
        return round(word_count / 200.0, 2)
    
    def _calculate_totals(self, results: List[TextMetrics]) -> Dict[str, Any]:
        """Calculate aggregate metrics across all files"""
        if not results:
            return {}
        
        total_words = sum(m.word_count for m in results)
        total_sentences = sum(m.sentence_count for m in results)
        total_paragraphs = sum(m.paragraph_count for m in results)
        total_chars = sum(m.character_count for m in results)
        total_size = sum(m.file_size_bytes for m in results)
        total_pages = sum(m.page_count for m in results if m.page_count is not None)
        
        # Aggregate keywords across all files
        all_keywords = {}
        for m in results:
            for word, count in m.top_keywords:
                all_keywords[word] = all_keywords.get(word, 0) + count
        
        top_keywords = sorted(all_keywords.items(), key=lambda x: x[1], reverse=True)[:15]
        
        return {
            'total_files': len(results),
            'total_words': total_words,
            'total_sentences': total_sentences,
            'total_paragraphs': total_paragraphs,
            'total_characters': total_chars,
            'total_size_bytes': total_size,
            'total_pages': total_pages if total_pages > 0 else None,
            'total_reading_time_minutes': round(total_words / 200.0, 2),
            'avg_words_per_file': round(total_words / len(results), 2),
            'avg_lexical_diversity': round(sum(m.lexical_diversity for m in results) / len(results), 4),
            'file_types': {
                'pdf': sum(1 for m in results if m.file_type == 'pdf'),
                'docx': sum(1 for m in results if m.file_type == 'docx'),
                'txt': sum(1 for m in results if m.file_type == 'txt')
            },
            'top_keywords_overall': top_keywords
        }


if __name__ == "__main__":
    # Example usage
    analyzer = TextAnalyzer()
    
    # Test with a single file
    print("=== Single File Analysis ===")
    try:
        metrics = analyzer.analyze_file("sample.txt")
        print(metrics.to_dict())
    except Exception as e:
        print(f"Error: {e}")
