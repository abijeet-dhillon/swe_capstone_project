"""
Minimal test suite for TextAnalyzer
Tests core functionality with TXT, DOCX, and PDF files
"""

import pytest
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from src.analyze.text_analyzer import TextAnalyzer, TextMetrics


@pytest.fixture
def test_files_dir(tmp_path):
    """Create temporary test files"""
    
    # Create TXT file
    txt_file = tmp_path / "test.txt"
    txt_content = """
Test Document

This is a test document for the text analyzer.
It contains multiple sentences. Each sentence is separated by periods.

This is a second paragraph with more content.
The analyzer should count words, sentences, and paragraphs correctly.
"""
    txt_file.write_text(txt_content)
    
    # Create DOCX file
    docx_file = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading('Test Heading 1', level=1)
    doc.add_paragraph('This is a test paragraph in a Word document.')
    doc.add_heading('Test Heading 2', level=2)
    doc.add_paragraph('Another paragraph with some test content.')
    doc.save(str(docx_file))
    
    # Create PDF file
    pdf_file = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_file), pagesize=letter)
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 730, "This is a test PDF with multiple lines.")
    c.drawString(100, 710, "It should be parsed correctly.")
    c.showPage()
    c.drawString(100, 750, "Second page content.")
    c.save()
    
    return {
        'txt': txt_file,
        'docx': docx_file,
        'pdf': pdf_file
    }


def test_analyze_txt_file(test_files_dir):
    """Test analyzing a TXT file"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['txt']))
    
    assert isinstance(metrics, TextMetrics)
    assert metrics.file_type == 'txt'
    assert metrics.word_count > 0
    assert metrics.sentence_count > 0
    assert metrics.paragraph_count > 0
    assert metrics.estimated_reading_time_minutes > 0
    assert 0 <= metrics.lexical_diversity <= 1
    assert len(metrics.top_keywords) > 0


def test_analyze_docx_file(test_files_dir):
    """Test analyzing a DOCX file"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['docx']))
    
    assert isinstance(metrics, TextMetrics)
    assert metrics.file_type == 'docx'
    assert metrics.word_count > 0
    assert metrics.heading_count is not None
    assert metrics.heading_count >= 2  # We added 2 headings
    assert metrics.heading_breakdown is not None


def test_analyze_pdf_file(test_files_dir):
    """Test analyzing a PDF file"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['pdf']))
    
    assert isinstance(metrics, TextMetrics)
    assert metrics.file_type == 'pdf'
    assert metrics.word_count > 0
    assert metrics.page_count is not None
    assert metrics.page_count == 2  # We created 2 pages


def test_batch_analysis(test_files_dir):
    """Test batch analysis with multiple files"""
    analyzer = TextAnalyzer()
    file_paths = [
        str(test_files_dir['txt']),
        str(test_files_dir['docx']),
        str(test_files_dir['pdf'])
    ]
    
    results = analyzer.analyze_batch(file_paths)
    
    # Check structure
    assert 'files' in results
    assert 'totals' in results
    assert len(results['files']) == 3
    
    # Check totals
    totals = results['totals']
    assert totals['total_files'] == 3
    assert totals['total_words'] > 0
    assert totals['total_sentences'] > 0
    assert totals['total_reading_time_minutes'] > 0
    assert 'file_types' in totals
    assert totals['file_types']['txt'] == 1
    assert totals['file_types']['docx'] == 1
    assert totals['file_types']['pdf'] == 1


def test_to_dict_conversion(test_files_dir):
    """Test that metrics can be converted to dictionary"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['txt']))
    
    result_dict = metrics.to_dict()
    
    assert isinstance(result_dict, dict)
    assert 'file_name' in result_dict
    assert 'word_count' in result_dict
    assert 'top_keywords' in result_dict
    assert result_dict['file_type'] == 'txt'


def test_file_not_found():
    """Test error handling for non-existent file"""
    analyzer = TextAnalyzer()
    
    with pytest.raises(FileNotFoundError):
        analyzer.analyze_file("nonexistent_file.txt")


def test_unsupported_file_type(tmp_path):
    """Test error handling for unsupported file type"""
    analyzer = TextAnalyzer()
    
    # Create a file with unsupported extension
    unsupported_file = tmp_path / "test.xyz"
    unsupported_file.write_text("test content")
    
    with pytest.raises(ValueError, match="Unsupported file type"):
        analyzer.analyze_file(str(unsupported_file))


def test_keyword_extraction(test_files_dir):
    """Test that keywords are extracted correctly"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['txt']))
    
    # Check keywords format
    assert isinstance(metrics.top_keywords, list)
    assert len(metrics.top_keywords) > 0
    
    # Each keyword should be a tuple of (word, count)
    for keyword, count in metrics.top_keywords:
        assert isinstance(keyword, str)
        assert isinstance(count, int)
        assert count > 0


def test_reading_time_calculation(test_files_dir):
    """Test reading time estimation"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['txt']))
    
    # Reading time should be word_count / 200
    expected_time = metrics.word_count / 200.0
    assert abs(metrics.estimated_reading_time_minutes - expected_time) < 0.01


def test_lexical_diversity(test_files_dir):
    """Test lexical diversity calculation"""
    analyzer = TextAnalyzer()
    metrics = analyzer.analyze_file(str(test_files_dir['txt']))
    
    # Lexical diversity should be between 0 and 1
    assert 0 <= metrics.lexical_diversity <= 1
    
    # For normal text, it should be reasonable (not 0 or 1)
    assert 0.1 <= metrics.lexical_diversity <= 0.9
