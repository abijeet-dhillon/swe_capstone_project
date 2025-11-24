#!/usr/bin/env python3
"""
Example Text Analysis
Creates sample files (TXT, DOCX, MD, PDF) and analyzes them
Run: python3 src/analyze/example_txt_analysis.py
"""

import json
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from text_analyzer import TextAnalyzer


def create_sample_files():
    """Create sample files for testing"""
    
    # Create TXT file
    txt_content = """
Digital Work Artifact Mining

This project analyzes digital work artifacts from personal computers.
The system extracts meaningful insights from code repositories, documents, and media files.

Key Features:
- Privacy-first design with user consent
- Support for multiple file formats (PDF, DOCX, TXT)
- Statistical analysis without requiring LLM services
- Batch processing capabilities for multiple files

The text analyzer component processes various document types.
It extracts metrics like word count, reading time, and keyword frequency.
These metrics help users showcase their productivity and skill development.

Technical Implementation:
The analyzer uses pdfminer for PDF extraction, python-docx for Word documents,
and built-in Python libraries for text files. All processing is done locally
to ensure user privacy and data security.
"""
    Path("sample_document.txt").write_text(txt_content)
    print("✅ Created sample_document.txt")
    
    # Create MD file
    md_content = """
# Project README

## Overview
This is a sample markdown document for the text analyzer demonstration.

## Technologies Used
- Python 3.11+
- FastAPI for API development
- SQLAlchemy for database operations
- OpenAI API for LLM integration

## Features
- User authentication and authorization
- Real-time data processing
- Comprehensive analytics dashboard
- Export functionality to multiple formats

## Installation
The project uses standard Python package management.
All dependencies are listed in requirements.txt for easy installation.

## Usage
The system provides both command-line and programmatic interfaces.
Users can analyze individual files or process entire directories in batch mode.
"""
    Path("sample_readme.md").write_text(md_content)
    print("✅ Created sample_readme.md")
    
    # Create DOCX file
    doc = Document()
    doc.add_heading('Technical Report', 0)
    doc.add_paragraph('This document demonstrates the text analyzer capabilities.')
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'The text analyzer extracts statistical metrics from documents. '
        'It supports PDF, DOCX, and TXT formats without requiring external AI services.'
    )
    doc.add_heading('Methodology', level=1)
    doc.add_paragraph(
        'The analyzer uses document parsing libraries to extract text content. '
        'Statistical analysis is performed using Python built-in functions and regex. '
        'Keyword extraction uses frequency-based algorithms with stop word filtering.'
    )
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        'The system successfully extracts word counts, sentence counts, and lexical diversity. '
        'Reading time estimates are calculated based on average reading speed. '
        'Top keywords are identified through frequency analysis.'
    )
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The text analyzer provides comprehensive metrics for document analysis. '
        'It enables users to understand their work artifacts quantitatively.'
    )
    doc.save("sample_report.docx")
    print("✅ Created sample_report.docx")
    
    # Create PDF file
    pdf_path = "sample_presentation.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    
    # Page 1
    c.setFont("Helvetica-Bold", 24)
    c.drawString(100, height - 100, "Digital Work Artifacts")
    c.setFont("Helvetica", 12)
    y = height - 150
    lines = [
        "This presentation covers the text analysis component.",
        "The system processes documents and extracts key metrics.",
        "",
        "Key Capabilities:",
        "- Multi-format support (PDF, DOCX, TXT)",
        "- Statistical analysis and keyword extraction",
        "- Batch processing for multiple files",
        "- Privacy-first local processing",
        "",
        "The analyzer provides quantitative insights into document content.",
        "Metrics include word count, reading time, and lexical diversity.",
    ]
    for line in lines:
        c.drawString(100, y, line)
        y -= 20
    
    # Page 2
    c.showPage()
    c.setFont("Helvetica-Bold", 18)
    c.drawString(100, height - 100, "Technical Details")
    c.setFont("Helvetica", 12)
    y = height - 150
    lines2 = [
        "Implementation:",
        "The text analyzer is built with Python and uses specialized libraries.",
        "PDF parsing: pdfminer.six",
        "DOCX parsing: python-docx",
        "Text processing: built-in Python libraries",
        "",
        "Output Format:",
        "Results are returned as structured dictionaries with comprehensive metrics.",
        "Both individual file analysis and batch processing are supported.",
        "Aggregate statistics are calculated for multiple file analysis.",
    ]
    for line in lines2:
        c.drawString(100, y, line)
        y -= 20
    
    c.save()
    print("✅ Created sample_presentation.pdf")


def analyze_files():
    """Analyze the created sample files"""
    
    files = [
        "sample_document.txt",
        "sample_readme.md",
        "sample_report.docx",
        "sample_presentation.pdf"
    ]
    
    analyzer = TextAnalyzer()
    
    print("\n" + "="*70)
    print("ANALYZING SAMPLE FILES")
    print("="*70)
    
    # Analyze in batch
    results = analyzer.analyze_batch(files)
    
    # Print results
    print("\n" + "="*70)
    print("BATCH ANALYSIS RESULTS (Dictionary)")
    print("="*70)
    print(json.dumps(results, indent=2))


def cleanup_files():
    """Remove created sample files"""
    files = [
        "sample_document.txt",
        "sample_readme.md",
        "sample_report.docx",
        "sample_presentation.pdf"
    ]
    
    print("\n" + "="*70)
    print("CLEANUP")
    print("="*70)
    for file in files:
        if Path(file).exists():
            Path(file).unlink()
            print(f"🗑️  Removed {file}")


def main():
    """Main execution"""
    print("\n" + "="*70)
    print("TEXT ANALYZER - EXAMPLE DEMONSTRATION")
    print("="*70)
    
    # Create sample files
    print("\n📝 Creating sample files...")
    create_sample_files()
    
    # Analyze them
    analyze_files()
    
    # Cleanup
    cleanup_files()
    
    print("\n" + "="*70)
    print("✅ EXAMPLE COMPLETE!")
    print("="*70)
    print("\nThe text analyzer successfully processed all file types.")
    print("You can now use it with your own files:")
    print("  python3 analyze_text.py your_document.pdf")
    print()


if __name__ == "__main__":
    main()
